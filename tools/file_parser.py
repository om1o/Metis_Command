"""
File parser — universal text extraction for the Scholar agent.

Supported:
    .pdf          via pdfplumber (already pulled transitively), pypdfium2 fallback
    .docx         via python-docx (lazy import; optional)
    .xlsx, .xls   via openpyxl / pandas
    .csv          via csv / pandas
    .md, .txt     raw
    .html, .htm   BeautifulSoup fallback, plain text strip otherwise
    .json         pretty-print
    images        OCR via Tesseract if available, else llava description

Every file runs through a size gate and returns a normalized dict:
    {"ok": bool, "path": str, "kind": str, "text": str, "pages": int | None}
"""

from __future__ import annotations

import csv
import json
import re
from pathlib import Path
from typing import Any

from safety import audited, require_safe_path


MAX_CHARS = 120_000


@audited("file_parser.parse")
def parse(path: str) -> dict[str, Any]:
    p = require_safe_path(path)
    if not p.exists() or not p.is_file():
        return {"ok": False, "error": f"not found: {p}"}

    ext = p.suffix.lower().lstrip(".")
    try:
        if ext == "pdf":
            return _parse_pdf(p)
        if ext == "docx":
            return _parse_docx(p)
        if ext in {"xlsx", "xls"}:
            return _parse_excel(p)
        if ext == "csv":
            return _parse_csv(p)
        if ext == "json":
            return _parse_json(p)
        if ext in {"html", "htm"}:
            return _parse_html(p)
        if ext in {"png", "jpg", "jpeg", "webp", "bmp"}:
            return _parse_image(p)
        return _parse_text(p)
    except Exception as e:
        return {"ok": False, "error": str(e), "path": str(p)}


# ── Individual parsers ───────────────────────────────────────────────────────

def _parse_text(p: Path) -> dict[str, Any]:
    text = p.read_text(encoding="utf-8", errors="replace")[:MAX_CHARS]
    return {"ok": True, "path": str(p), "kind": "text", "text": text, "pages": None}


def _parse_pdf(p: Path) -> dict[str, Any]:
    text_parts: list[str] = []
    pages = 0
    try:
        import pdfplumber  # type: ignore
        with pdfplumber.open(str(p)) as pdf:
            for page in pdf.pages:
                pages += 1
                text_parts.append(page.extract_text() or "")
    except Exception:
        try:
            import pypdfium2 as pdfium  # type: ignore
            doc = pdfium.PdfDocument(str(p))
            for i in range(len(doc)):
                pages += 1
                text_parts.append(doc[i].get_textpage().get_text_range())
        except Exception as e:
            return {"ok": False, "error": f"pdf parse failed: {e}", "path": str(p)}
    text = "\n\n".join(text_parts)[:MAX_CHARS]
    return {"ok": True, "path": str(p), "kind": "pdf", "text": text, "pages": pages}


def _parse_docx(p: Path) -> dict[str, Any]:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return {"ok": False, "error": "python-docx not installed", "path": str(p)}
    doc = Document(str(p))
    text = "\n".join(par.text for par in doc.paragraphs)[:MAX_CHARS]
    return {"ok": True, "path": str(p), "kind": "docx", "text": text, "pages": None}


def _parse_excel(p: Path) -> dict[str, Any]:
    try:
        import openpyxl  # type: ignore
        wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
    except Exception as e:
        return {"ok": False, "error": f"xlsx open failed: {e}", "path": str(p)}
    chunks: list[str] = []
    for sheet in wb.worksheets:
        chunks.append(f"### Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True, max_rows=400):
            chunks.append(" | ".join("" if v is None else str(v) for v in row))
    text = "\n".join(chunks)[:MAX_CHARS]
    return {"ok": True, "path": str(p), "kind": "xlsx", "text": text, "pages": len(wb.worksheets)}


def _parse_csv(p: Path) -> dict[str, Any]:
    rows: list[str] = []
    with p.open("r", encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.reader(f)
        for i, row in enumerate(reader):
            if i > 2000:
                rows.append("… (truncated)")
                break
            rows.append(" | ".join(row))
    return {"ok": True, "path": str(p), "kind": "csv", "text": "\n".join(rows)[:MAX_CHARS], "pages": None}


def _parse_json(p: Path) -> dict[str, Any]:
    try:
        data = json.loads(p.read_text(encoding="utf-8", errors="replace"))
        text = json.dumps(data, indent=2, ensure_ascii=False)[:MAX_CHARS]
    except Exception:
        text = p.read_text(encoding="utf-8", errors="replace")[:MAX_CHARS]
    return {"ok": True, "path": str(p), "kind": "json", "text": text, "pages": None}


def _parse_html(p: Path) -> dict[str, Any]:
    raw = p.read_text(encoding="utf-8", errors="replace")
    try:
        from bs4 import BeautifulSoup  # type: ignore
        text = BeautifulSoup(raw, "html.parser").get_text("\n")
    except Exception:
        text = re.sub(r"<[^>]+>", " ", raw)
        text = re.sub(r"\s+", " ", text)
    return {"ok": True, "path": str(p), "kind": "html", "text": text[:MAX_CHARS], "pages": None}


def _parse_image(p: Path) -> dict[str, Any]:
    """Try Tesseract OCR, then fall back to asking llava what's in the image."""
    try:
        import pytesseract  # type: ignore
        from PIL import Image
        text = pytesseract.image_to_string(Image.open(str(p)))
        if text.strip():
            return {"ok": True, "path": str(p), "kind": "image-ocr", "text": text[:MAX_CHARS], "pages": None}
    except Exception:
        pass

    try:
        import base64
        from brain_engine import chat_by_role
        img_b64 = base64.b64encode(p.read_bytes()).decode("ascii")
        reply = chat_by_role(
            "vision",
            [{"role": "user", "content": "Describe this image. Extract any visible text verbatim.",
              "images": [img_b64]}],
        )
        return {"ok": True, "path": str(p), "kind": "image-vision", "text": reply[:MAX_CHARS], "pages": None}
    except Exception as e:
        return {"ok": False, "error": f"image parse failed: {e}", "path": str(p)}


# ── CrewAI adapter ───────────────────────────────────────────────────────────

def as_crewai_tool():
    try:
        from crewai.tools import tool  # type: ignore
    except Exception:
        return None

    @tool("ParseFile")
    def _parse(path: str) -> str:
        """Extract text from a PDF/docx/xlsx/csv/html/image/text file."""
        import json as _json
        return _json.dumps(parse(path))

    return _parse
